# -*- coding: utf-8 -*-
import os
import re
import json
import warnings
import httpx
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

warnings.filterwarnings("ignore")
load_dotenv(override=True)

MODEL_CHAT = os.getenv("MODEL_CHAT", "glm-5-outside")


def _get_client():
    return OpenAI(
        api_key=os.getenv("AIRCHINA_API_KEY"),
        base_url=os.getenv("AIRCHINA_BASE_URL"),
        http_client=httpx.Client(verify=False),
    )


def _parse_json(raw):
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


_EXTRACT_PROMPT_TMPL = """你是一个公文处理专家。请从以下呈批件文本中提取关键信息，以JSON格式返回。

要求：
- 文件编号：如"中航集团资产呈〔2024〕231号"，找不到填null
- 项目名称：项目完整名称，找不到填null
- 项目概况：含规模、面积、栋数、投资情况等，尽量完整，找不到填null
- 总投资金额：如"约X.XX亿元"或"XXX万元"，找不到填null
- 被授权单位：承办该项目的单位/部门名称，找不到填null
- 授权事项：列表形式，从呈批件中提炼所有需上级授权的事项，每条30字以内
- 拟稿部门：起草本文件的部门，找不到填null
- 文件日期：文件落款日期，格式YYYY年MM月DD日，找不到填null
- 主送机构：如"党组"、"总经理办公会"、"董事会"等，找不到填null
- 授权单位：需出具授权书的单位全称（即下级向上级授权时，下级单位的全称），找不到填null
- 被授权上级单位：接受授权的上级单位全称，找不到填null
- 转授权单位：被授权单位可再转委托的第三方单位，若无则填null
- 授权期限：授权有效期的完整描述，直接摘录原文表述，找不到填null
- 授权份数：授权书需办理的份数（整数），找不到填null
- 印章要求：办理授权需使用的印章类型，直接摘录原文，找不到填null
- 授权背景：从呈批件中综合提炼为何需要进行授权、授权的必要性和具体安排，1~2句话，直接综合原文内容，不编造，不套用固定句式，找不到填null

只返回JSON，不要其他文字：
{{"文件编号":null,"项目名称":null,"项目概况":null,"总投资金额":null,"被授权单位":null,"授权事项":[],"拟稿部门":null,"文件日期":null,"主送机构":null,"授权单位":null,"被授权上级单位":null,"转授权单位":null,"授权期限":null,"授权份数":null,"印章要求":null,"授权背景":null}}

呈批件文本（前10000字）：
{text}"""


_DRAFT_PROMPT_TMPL = """你是一名精通国有企业公文写作的专家。请根据以下信息起草一份正式的授权请示。

输入信息：
- 呈批件文件编号：{wh}
- 呈批件标题（主送机构请示内容）：{xm}
- 授权背景（从呈批件提炼的授权原因和安排）：{bg}
- 授权单位（出具授权书的单位）：{qdw}
- 被授权单位：{bqdw}
- 转授权单位（若有）：{zqdw}
- 授权事项：
{sq}
- 授权期限：{qx}
- 授权份数：{fs}
- 印章要求：{yz}
- 拟稿部门：{bm}
- 文件日期：{rq}

按以下格式生成，只输出文档正文，不要任何解释，不要Markdown符号，不要"标题："前缀：

【来文依据段】
根据《{xm}》（{wh}）{bg}

【授权说明段】
现拟将{qdw_short}负责开展相关工作（具体范围请补充）授权给{bqdw_short}，需由{qdw_short}出具授权委托书。

【授权内容段】
具体的授权内容包括：将上方授权事项展开为完整规范的表述，各条用分号连接，最后一条用句号结束。

【实务条款】（每项单独一行，对应字段为null则跳过该行）
授权委托期限：{qx}。
授权办理份数：{fs}份。
办理授权时需使用{qdw_short}公章及法定代表人签字章。
妥否，请批示。

【附件清单】
附件：1. 授权书（{qdw_short}授权{bqdw_short}）
{fj_extra}

【落款】
{bm}
{rq}"""


_AUTH_LETTER_PROMPT_TMPL = """你是一名精通国有企业公文写作的专家。请根据以下信息，生成授权书中"授权委托事项范围"段落的完整内容。

授权单位：{qdw}
被授权单位：{bqdw}
授权事项列表：
{sq}

要求：
- 将上述授权事项展开为完整规范的授权委托事项范围描述
- 描述完整、边界清晰、措辞正式
- 直接输出段落文字，不要标题，不要编号，不要Markdown

只输出"授权委托事项范围"的正文内容："""


def extract_approval_info(pdf_text):
    prompt = _EXTRACT_PROMPT_TMPL.replace("{text}", pdf_text[:10000])
    resp = _get_client().chat.completions.create(
        model=MODEL_CHAT,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=3000,
        temperature=0.1,
    )
    try:
        return _parse_json(resp.choices[0].message.content)
    except Exception:
        return {
            "文件编号": None, "项目名称": None, "项目概况": None,
            "总投资金额": None, "被授权单位": None, "授权事项": [],
            "拟稿部门": None, "文件日期": None, "主送机构": None,
            "授权单位": None, "被授权上级单位": None, "转授权单位": None,
            "授权期限": None, "授权份数": None, "印章要求": None, "授权背景": None,
        }


def _short_name(full_name):
    """提取括号内简称，否则取前10字。"""
    if not full_name:
        return full_name or "（待填写）"
    m = re.search(r'[（(]简称(.+?)[）)]', full_name)
    if m:
        return m.group(1)
    return full_name[:10] if len(full_name) > 10 else full_name


def draft_auth_request(info):
    items = info.get("授权事项") or []
    items_text = "\n".join(
        "{}. {}".format(i + 1, item) for i, item in enumerate(items)
    ) if items else "（待补充）"

    qdw = info.get("授权单位") or "（待填写）"
    bqdw = info.get("被授权上级单位") or info.get("被授权单位") or "（待填写）"
    zqdw = info.get("转授权单位")
    qdw_short = _short_name(qdw)
    bqdw_short = _short_name(bqdw)

    # 附件额外行
    fj_extra = ""
    if zqdw:
        zqdw_short = _short_name(zqdw)
        fj_extra = "      2. 集团公司授权书（{}授权{}）".format(bqdw_short, zqdw_short)

    prompt = (
        _DRAFT_PROMPT_TMPL
        .replace("{wh}", info.get("文件编号") or "（未提取到）")
        .replace("{xm}", info.get("项目名称") or "（未提取到）")
        .replace("{bg}", info.get("授权背景") or "根据上述文件要求，需办理相关授权手续")
        .replace("{qdw}", qdw)
        .replace("{bqdw}", bqdw)
        .replace("{zqdw}", zqdw or "无")
        .replace("{sq}", items_text)
        .replace("{qx}", info.get("授权期限") or "（待填写）")
        .replace("{fs}", str(info.get("授权份数")) if info.get("授权份数") else "（待填写）")
        .replace("{yz}", info.get("印章要求") or "（待填写）")
        .replace("{bm}", info.get("拟稿部门") or "（拟稿部门）")
        .replace("{rq}", info.get("文件日期") or "（日期）")
        .replace("{qdw_short}", qdw_short)
        .replace("{bqdw_short}", bqdw_short)
        .replace("{fj_extra}", fj_extra)
    )

    resp = _get_client().chat.completions.create(
        model=MODEL_CHAT,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=4000,
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()


def draft_auth_letter(info):
    """生成授权书正文（框架固定，AI只生成授权委托事项范围段落）。"""
    items = info.get("授权事项") or []
    items_text = "\n".join(
        "{}. {}".format(i + 1, item) for i, item in enumerate(items)
    ) if items else "（待补充）"

    qdw = info.get("授权单位") or "（待填写）"
    bqdw = info.get("被授权上级单位") or info.get("被授权单位") or "（待填写）"
    qdw_short = _short_name(qdw)

    # 用AI生成授权委托事项范围段落
    scope_prompt = (
        _AUTH_LETTER_PROMPT_TMPL
        .replace("{qdw}", qdw)
        .replace("{bqdw}", bqdw)
        .replace("{sq}", items_text)
    )
    resp = _get_client().chat.completions.create(
        model=MODEL_CHAT,
        messages=[{"role": "user", "content": scope_prompt}],
        max_tokens=2000,
        temperature=0.2,
    )
    scope_text = resp.choices[0].message.content.strip()

    # 年份
    date_str = info.get("文件日期") or ""
    year = date_str[:4] if date_str and len(date_str) >= 4 else "    "

    letter = (
        "法定代表人授权书\n\n"
        "{qdw_short}授权字（{year}）第      号\n\n"
        "授权单位：{qdw}\n"
        "注册地址：\n"
        "法定代表人：                ，职务：\n\n"
        "被授权单位：{bqdw}\n"
        "注册地址：\n"
        "法定代表人：                ，职务：\n\n"
        "授权委托事项范围：{scope}\n\n"
        "被授权单位权限：被授权单位应在授权委托事项范围内依法合规地行使权力、履行义务。"
        "被授权单位在授权委托事项范围内代表授权单位所签署的法律文件具有法律效力，"
        "与授权单位的行为具有同等法律效力，授权单位予以认可。"
        "授权单位同意，被授权单位有转委托权。\n\n"
        "授权期限：{qx}。如{qdw}终止对被授权人的授权后，本授权书即失效。\n\n"
        "授权单位印签：\n"
        "法定代表人：\n"
        "                                            {rq}"
    ).format(
        qdw_short=qdw_short,
        year=year,
        qdw=qdw,
        bqdw=bqdw,
        scope=scope_text,
        qx=info.get("授权期限") or "（待填写）",
        rq=info.get("文件日期") or "（日期）",
    )
    return letter


_AUTH_LEDGER_HEADERS = [
    "序号", "编号", "经办人", "授权人", "代理人", "印章", "份数",
    "授权起始日期", "授权终止日期", "授权内容概要", "办理时间",
    "代理人签字版是否发送回来", "文号", "责任者", "题目", "归档日期", "页数",
]


def init_auth_ledger(ledger_path):
    """若台账文件不存在则自动创建（含表头）。"""
    import os
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    os.makedirs(os.path.dirname(ledger_path), exist_ok=True)
    if os.path.exists(ledger_path):
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "授权委托台账"

    for col, header in enumerate(_AUTH_LEDGER_HEADERS, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(fill_type="solid", fgColor="4472C4")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    col_widths = [6, 12, 10, 22, 22, 18, 6, 14, 14, 40, 14, 18, 16, 12, 40, 12, 6]
    for col, width in enumerate(col_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width
    ws.row_dimensions[1].height = 22

    wb.save(ledger_path)


def record_to_ledger(info, title, ledger_path):
    """向台账 Excel 追加一行，失败静默处理。"""
    try:
        import openpyxl
        from openpyxl.styles import Alignment as XlAlign
        init_auth_ledger(ledger_path)
        wb = openpyxl.load_workbook(ledger_path)
        ws = wb.active

        # 当前最大序号
        max_seq = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is not None:
                try:
                    max_seq = max(max_seq, int(row[0]))
                except (TypeError, ValueError):
                    pass

        # 授权内容概要
        items = info.get("授权事项") or []
        summary = "；".join(items)[:200] if items else ""

        new_row = [
            max_seq + 1,           # 序号
            None,                   # 编号（人工填写）
            info.get("拟稿人"),     # 经办人
            info.get("授权单位"),   # 授权人
            info.get("被授权上级单位") or info.get("被授权单位"),  # 代理人
            info.get("印章要求"),   # 印章
            info.get("授权份数"),   # 份数
            info.get("文件日期"),   # 授权起始日期
            info.get("授权期限"),   # 授权终止日期
            summary,                # 授权内容概要
            info.get("文件日期"),   # 办理时间
            None,                   # 代理人签字版是否发送回来
            info.get("文件编号"),   # 文号
            None,                   # 责任者
            title,                  # 题目
            None,                   # 归档日期
            None,                   # 页数
        ]
        next_row = ws.max_row + 1
        for col, value in enumerate(new_row, start=1):
            cell = ws.cell(row=next_row, column=col, value=value)
            cell.alignment = XlAlign(vertical="center")
            if next_row % 2 == 0:
                from openpyxl.styles import PatternFill
                cell.fill = PatternFill(fill_type="solid", fgColor="DCE6F1")
        wb.save(ledger_path)
        return True
    except Exception:
        return False


def _set_run_font(run, cn_font, size_pt):
    """设置字体：中文用 cn_font，西文用 Times New Roman。"""
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),   'Times New Roman')
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'),       cn_font)


def _build_footer_para(para, alignment):
    """在指定段落插入 -页码- 格式的 PAGE 域。"""
    para.alignment = alignment

    def _text_run(p, text):
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p._p.append(r)

    def _field_run(p, fld_type):
        r = OxmlElement('w:r')
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), fld_type)
        r.append(fc)
        p._p.append(r)

    def _instr_run(p, instr):
        r = OxmlElement('w:r')
        it = OxmlElement('w:instrText')
        it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        it.text = instr
        r.append(it)
        p._p.append(r)

    _text_run(para, '-')
    _field_run(para, 'begin')
    _instr_run(para, ' PAGE \\* MERGEFORMAT ')
    _field_run(para, 'separate')
    _text_run(para, '1')
    _field_run(para, 'end')
    _text_run(para, '-')


def _setup_page_numbers(doc, section):
    """奇数页右下、偶数页左下，格式 -n-。"""
    settings_elem = doc.settings.element
    if settings_elem.find(qn('w:evenAndOddHeaders')) is None:
        even_odd = OxmlElement('w:evenAndOddHeaders')
        settings_elem.insert(0, even_odd)

    odd_footer = section.footer
    _build_footer_para(odd_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)

    even_footer = section.even_page_footer
    _build_footer_para(even_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)


def save_as_docx(content, output_path):
    """授权请示 Word 文档（方正小标宋/黑体/仿宋格式）。"""
    doc = Document()

    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)
        section.footer_distance = Cm(1.75)
        _setup_page_numbers(doc, section)

    for line in content.splitlines():
        line_s = line.strip()
        if not line_s:
            doc.add_paragraph("")
            continue

        # 标题行（标题：xxx）
        if line_s.startswith("标题："):
            title_text = line_s[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            _set_run_font(run, '方正小标宋简体', 18)

        # 一级标题：一、二、三、…
        elif re.match(r'^[一二三四五六七八九十]+、', line_s):
            p = doc.add_paragraph()
            run = p.add_run(line_s)
            _set_run_font(run, '黑体', 16)

        # 二级标题：（一）（二）…
        elif re.match(r'^（[一二三四五六七八九十]+）', line_s):
            p = doc.add_paragraph()
            run = p.add_run(line_s)
            _set_run_font(run, '楷体', 16)

        # 附件行
        elif line_s.startswith("附件：") or re.match(r'^附件\s*\d+', line_s):
            p = doc.add_paragraph()
            run = p.add_run(line_s)
            _set_run_font(run, '仿宋_GB2312', 16)

        # 普通正文
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(21)
            run = p.add_run(line_s)
            _set_run_font(run, '仿宋_GB2312', 16)

    doc.save(output_path)


def save_auth_letter_as_docx(content, output_path):
    """授权书 Word 文档格式。"""
    doc = Document()

    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    lines = content.splitlines()
    for i, line in enumerate(lines):
        line_s = line.strip()
        if not line_s:
            doc.add_paragraph("")
            continue

        # 主标题
        if line_s == "法定代表人授权书":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_s)
            _set_run_font(run, '方正小标宋简体', 18)

        # 编号行（含"授权字"）
        elif "授权字" in line_s:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line_s)
            _set_run_font(run, '仿宋_GB2312', 14)

        # 字段标签行（授权单位：/被授权单位：/授权委托事项范围：等）
        elif re.match(r'^(授权单位|被授权单位|注册地址|法定代表人|授权委托事项范围|被授权单位权限|授权期限|授权单位印签)[:：]', line_s):
            p = doc.add_paragraph()
            run = p.add_run(line_s)
            _set_run_font(run, '黑体', 14)

        # 签字区（纯空白 + 日期结尾）
        elif line_s.startswith("法定代表人：") and i > 5:
            p = doc.add_paragraph()
            run = p.add_run(line_s)
            _set_run_font(run, '仿宋_GB2312', 14)

        elif re.match(r'^\s+\d{4}年', line_s):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line_s.strip())
            _set_run_font(run, '仿宋_GB2312', 14)

        # 正文段落
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(28)
            run = p.add_run(line_s)
            _set_run_font(run, '仿宋_GB2312', 14)

    doc.save(output_path)
